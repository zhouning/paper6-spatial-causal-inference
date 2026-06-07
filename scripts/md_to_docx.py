"""Convert world-model-technical-report.md to a formatted Word document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_md(md_text: str) -> list[dict]:
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            quote_lines = []
            while i < len(lines) and (lines[i].startswith("> ") or lines[i].startswith(">")):
                quote_lines.append(lines[i].lstrip("> ").strip())
                i += 1
            blocks.append({"type": "blockquote", "text": " ".join(quote_lines)})
            continue

        # Code block
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            # Parse table
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl):
                    continue  # skip separator
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # Regular paragraph (collect consecutive non-empty lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") \
                and not lines[i].startswith("```") and not lines[i].startswith("> ") \
                and not ("|" in lines[i] and i + 1 < len(lines) and "|" in lines[i + 1] and "-" in lines[i + 1]):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            text = " ".join(para_lines)
            # Check if it's a list
            if para_lines[0].strip().startswith(("- ", "* ", "1. ")):
                for pl in para_lines:
                    blocks.append({"type": "list_item", "text": re.sub(r"^[\-\*\d+\.]\s+", "", pl.strip())})
            else:
                blocks.append({"type": "paragraph", "text": text})

    return blocks


def add_formatted_text(paragraph, text: str):
    """Add text with basic inline formatting (**bold**, *italic*, `code`)."""
    # Split by formatting markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            # Clean up markdown links [text](url) -> text
            cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", part)
            paragraph.add_run(cleaned)


def build_docx(blocks: list[dict], output_path: str):
    """Build Word document from parsed blocks."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Set CJK font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            heading = doc.add_heading(level=min(level, 4))
            add_formatted_text(heading, block["text"])
            # Style heading fonts
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                if run.font.name is None or run.font.name == "Calibri":
                    run.font.name = "Calibri"

        elif btype == "paragraph":
            para = doc.add_paragraph()
            add_formatted_text(para, block["text"])

        elif btype == "blockquote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(block["text"])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif btype == "code":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            # Add shading
            pPr = para._element.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:fill"): "F5F5F5",
                qn("w:val"): "clear",
            })
            pPr.append(shd)
            run = para.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        elif btype == "list_item":
            para = doc.add_paragraph(style="List Bullet")
            add_formatted_text(para, block["text"])

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols, style="Light Grid Accent 1")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < n_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_formatted_text(para, cell_text.strip())
                        para.paragraph_format.space_after = Pt(2)
                        for run in para.runs:
                            run.font.size = Pt(9)
                        # Bold header row
                        if ri == 0:
                            for run in para.runs:
                                run.bold = True

    doc.save(output_path)
    return output_path


import sys

def main():
    if len(sys.argv) > 2:
        md_path = Path(sys.argv[1])
        output_path = Path(sys.argv[2])
    else:
        repo_root = Path(__file__).resolve().parent.parent
        md_path = repo_root / "docs" / "background" / "world-model-technical-review.md"
        output_path = repo_root / "docs" / "background" / "generated" / "world-model-technical-review.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_md(md_text)
    result = build_docx(blocks, str(output_path))
    print(f"Word document generated: {result}")
    print(f"Blocks parsed: {len(blocks)}")


if __name__ == "__main__":
    main()
