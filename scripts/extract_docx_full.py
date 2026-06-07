#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Extract .docx full content including tables, images, and all paragraph details"""
import sys
from docx import Document
from docx.oxml.ns import qn

def extract_full_docx(filepath):
    doc = Document(filepath)
    content = []

    # Mapping of style IDs to heading levels
    heading_map = {}

    content.append("=" * 80)
    content.append("DOCUMENT STRUCTURE WITH FULL CONTENT")
    content.append("=" * 80)

    para_idx = 0
    table_idx = 0

    # Iterate through all body elements in order
    for element in doc.element.body:
        if element.tag == qn('w:p'):
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                style_name = para.style.name if para.style else "Normal"
                text = para.text.strip()

                # Determine heading level
                if 'Heading' in style_name or style_name.startswith('heading'):
                    level = ''.join(filter(str.isdigit, style_name)) or '?'
                    prefix = f"[H{level}]"
                elif '21bc9c4b' in style_name:
                    prefix = "[H1]"  # 一级标题
                elif '71e7dc79' in style_name:
                    prefix = "[H2]"  # 二级标题
                elif 'b63ee27f' in style_name:
                    prefix = "[H3]"  # 三级标题
                elif style_name == 'Title':
                    prefix = "[TITLE]"
                else:
                    prefix = f"[{style_name}]"

                if text:
                    content.append(f"{prefix} {text}")
                elif 'H' in prefix:
                    content.append(f"{prefix} (empty)")

            para_idx += 1

        elif element.tag == qn('w:tbl'):
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                content.append("")
                content.append(f"--- TABLE {table_idx + 1} ---")
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    content.append(f"  Row {row_idx}: {' | '.join(cells)}")
                content.append(f"--- END TABLE {table_idx + 1} ---")
                content.append("")
            table_idx += 1

    # Also list all images/shapes
    content.append("")
    content.append("=" * 80)
    content.append("DOCUMENT STATS")
    content.append(f"Total paragraphs: {len(doc.paragraphs)}")
    content.append(f"Total tables: {len(doc.tables)}")

    # Check for images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    content.append(f"Total images: {image_count}")
    content.append("=" * 80)

    return "\n".join(content)

if __name__ == "__main__":
    filepath = "时空数据中台产品详细设计V3.0.0.0（Data Agent部分）.docx"
    result = extract_full_docx(filepath)

    with open("doc_full_structure.txt", "w", encoding="utf-8") as f:
        f.write(result)

    print("Full extraction complete.")
